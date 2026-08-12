from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from docx import Document
import os

class ReportGenerator:
    """
    Generates Executive Audit Reports in PDF and Docx formats.
    """
    def __init__(self, output_dir="reports"):
        self.output_dir = output_dir
        os.makedirs(self.output_dir, exist_ok=True)

    def generate_pdf(self, audit_data, filename="Executive_Audit_Report.pdf"):
        path = os.path.join(self.output_dir, filename)
        c = canvas.Canvas(path, pagesize=letter)
        
        # Header
        c.setFont("Helvetica-Bold", 16)
        c.drawString(100, 750, "EXECUTIVE COMPLIANCE AUDIT REPORT")
        c.setFont("Helvetica", 12)
        c.drawString(100, 730, f"Developed by HSINI MOHAMED")
        c.line(100, 720, 500, 720)
        
        # Content
        c.drawString(100, 680, f"Framework: {audit_data.get('framework', 'ISO 27001')}")
        c.drawString(100, 660, f"Compliance Score: {audit_data.get('score', 0)}%")
        c.drawString(100, 640, f"Status: {audit_data.get('status', 'IN PROGRESS')}")
        
        c.setFont("Helvetica-Bold", 12)
        c.drawString(100, 600, "CONTROL GAP SUMMARY:")
        c.setFont("Helvetica", 10)
        y = 580
        for gap in audit_data.get('gaps', []):
            c.drawString(120, y, f"- {gap}")
            y -= 20
            
        c.save()
        return path

    def generate_docx(self, audit_data, filename="Executive_Audit_Report.docx"):
        path = os.path.join(self.output_dir, filename)
        doc = Document()
        
        doc.add_heading('EXECUTIVE COMPLIANCE AUDIT REPORT', 0)
        doc.add_paragraph(f"Developed by HSINI MOHAMED")
        
        doc.add_heading('Audit Summary', level=1)
        doc.add_paragraph(f"Framework: {audit_data.get('framework', 'ISO 27001')}")
        doc.add_paragraph(f"Overall Compliance Score: {audit_data.get('score', 0)}%")
        
        doc.add_heading('Identified Gaps', level=1)
        for gap in audit_data.get('gaps', []):
            doc.add_paragraph(gap, style='List Bullet')
            
        doc.save(path)
        return path
